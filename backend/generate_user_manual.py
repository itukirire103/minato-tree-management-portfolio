# -*- coding: utf-8 -*-
"""
generate_user_manual.py
--------------------------
(仮称)樹木管理システム ポートフォリオ用
簡易マニュアル(利用者向け、.docx)を生成する。

非機能要件#44(簡易マニュアル・説明会資料)に対応。
schema_manifest.json・create_security_roles.pyの内容を基に、
ロール別にできることを一覧化している。

実行方法:
    py generate_user_manual.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
import os

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_FILE = os.path.join(HERE, "簡易マニュアル.docx")

ROLES = [
    ("各所管理者", "全テーブルへのフルアクセス(組織全体)。樹木の登録・編集・削除、権限管理まで行える最上位の運用担当者"),
    ("区一般職員", "台帳の閲覧・登録・更新が中心。作業履歴・点検記録・苦情記録の管理を行う"),
    ("街路樹管理委託事業者", "担当エリアの樹木情報を閲覧・更新し、作業履歴を登録する"),
    ("協定管理者", "担当エリアの樹木情報・診断結果・点検記録を閲覧し、点検記録を登録する"),
    ("その他(閲覧専用)", "組織全体の情報を閲覧のみ行える(登録・更新は不可)"),
]


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("簡易マニュアル", level=0)
    sub = doc.add_paragraph("(仮称)樹木管理システム 構築プロジェクト")
    sub.runs[0].italic = True
    doc.add_paragraph("作成日: 2026-08-28")

    doc.add_heading("1. このシステムについて", level=1)
    doc.add_paragraph(
        "街路樹の位置・健全度・診断結果・作業履歴・苦情対応を一元管理するシステムです。"
        "ブラウザ(PC・スマートフォン・タブレット)からアクセスできます。追加のソフトウェアの"
        "インストールは不要です。"
    )

    doc.add_heading("2. ログイン方法", level=1)
    doc.add_paragraph("組織アカウント(メールアドレス)とパスワードでサインインします。", style="List Bullet")
    doc.add_paragraph("初回サインイン時、または一定期間ごとに多要素認証(MFA)の確認が求められます。", style="List Bullet")

    doc.add_heading("3. ロール別にできること", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "ロール"
    hdr[1].text = "できること"
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].paragraphs[0].runs[0].bold = True
    for name, desc in ROLES:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc

    doc.add_heading("4. 基本操作", level=1)

    doc.add_heading("4.1 樹木マスタを検索・閲覧する", level=2)
    doc.add_paragraph("左側メニューの「樹木マスタ」をクリックすると一覧が表示されます。", style="List Bullet")
    doc.add_paragraph("キーワード検索、または列見出しをクリックしての並べ替え・絞り込みができます。", style="List Bullet")

    doc.add_heading("4.2 樹木マップを使う", level=2)
    doc.add_paragraph("一覧画面左上のビュー切り替えから「樹木マップ」を選択すると、地図上にピンで表示されます。", style="List Bullet")
    doc.add_paragraph("ピンの色は健全度を表します(緑=良好、黄・橙=要注意、赤=要対応)。", style="List Bullet")
    doc.add_paragraph("ピンをクリックすると、その樹木の詳細情報が開きます。", style="List Bullet")
    doc.add_paragraph(
        "更新権限を持つ方は、ピンをドラッグして位置を修正できます。作成権限を持つ方は、"
        "地図の空いている場所をクリックして新しい樹木を登録できます(緯度・経度は自動入力されます)。",
        style="List Bullet",
    )

    doc.add_heading("4.3 作業履歴を登録する", level=2)
    doc.add_paragraph("該当する樹木の詳細画面から「作業履歴」タブを開き、新規登録します。", style="List Bullet")
    doc.add_paragraph(
        "作業前後の写真は、レコード右側の「タイムライン」欄から添付できます(📎アイコン、またはメモ欄への"
        "ドラッグ&ドロップ)。",
        style="List Bullet",
    )

    doc.add_heading("4.4 診断結果・点検記録を登録する", level=2)
    doc.add_paragraph("該当する樹木の詳細画面から、それぞれのタブで新規登録します。", style="List Bullet")
    doc.add_paragraph("診断カルテ(PDF)は、診断結果レコードの添付ファイル欄から登録します。", style="List Bullet")

    doc.add_heading("4.5 苦情・陳情を記録する", level=2)
    doc.add_paragraph("「苦情・陳情記録」から新規登録し、対応状況をステータスで管理します。", style="List Bullet")

    doc.add_heading("5. よくあるトラブル", level=1)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Light Grid Accent 1"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "症状"
    hdr2[1].text = "対処"
    hdr2[0].paragraphs[0].runs[0].bold = True
    hdr2[1].paragraphs[0].runs[0].bold = True
    faq = [
        ("画面が真っ白、または反応しない", "ブラウザの再読み込み(Ctrl+Shift+R)を試してください"),
        ("樹木マップが「地図の設定を読み込んでいます...」のまま止まる", "一度ブラウザタブを閉じて開き直してください。改善しない場合は運用担当者にご連絡ください"),
        ("編集・削除ボタンが表示されない", "ご自身のロールに権限が無い可能性があります。運用担当者にご確認ください"),
    ]
    for symptom, action in faq:
        row = table2.add_row().cells
        row[0].text = symptom
        row[1].text = action

    disclaimer = doc.add_paragraph(
        "本プロジェクトは東京都港区が公表した公募資料を参考にした個人の学習・ポートフォリオ目的の制作物であり、"
        "港区への正式な提案書・提出物ではありません。"
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.save(OUT_FILE)
    print(f"簡易マニュアル.docx を生成しました: {OUT_FILE}")


if __name__ == "__main__":
    main()
