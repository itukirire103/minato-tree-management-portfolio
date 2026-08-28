# -*- coding: utf-8 -*-
"""
generate_ops_manual.py
-------------------------
(仮称)樹木管理システム ポートフォリオ用
運用・障害対応手順書(.docx)を生成する。

非機能要件#9(運用手順書)・#10/#17(障害対応・再発防止)・#29(インシデント対応)に対応。

方針: 本システムはPower Platform(SaaS)上に構築されているため、インフラ・
パッチ適用等の大半はMicrosoft側が担う。運用者の役割は「監視」「異常時の
切り分け」「記録」に集約されるため、本手順書はその3点を中心に構成する。
本プロジェクトで実際に作成した検証スクリプト群を、そのまま定期点検ツール
として位置づけている点が特徴。

実行方法:
    py generate_ops_manual.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
import os

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_FILE = os.path.join(HERE, "運用_障害対応手順書.docx")


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("運用・障害対応手順書", level=0)
    sub = doc.add_paragraph("(仮称)樹木管理システム 構築プロジェクト")
    sub.runs[0].italic = True
    add_p(doc, "作成日: 2026-08-28")

    add_h(doc, "1. 本手順書の位置づけ")
    add_p(doc, (
        "本システムはMicrosoft Power Platform(SaaS)上に構築されているため、サーバーのパッチ適用・"
        "OSメンテナンス・物理的な障害対応はMicrosoft側が担う(非機能要件#26〜#30が該当)。"
        "運用者(本プロジェクトでは開発者自身)の役割は、以下の3点に集約される。"
    ))
    add_bullets(doc, [
        "日常的な健全性の監視(定期点検)",
        "異常発生時の切り分けと一次対応",
        "対応内容の記録・再発防止への反映",
    ])

    add_h(doc, "2. 日常運用: 定期点検")
    add_p(doc, (
        "本プロジェクトの構築ハーネスで作成した検証スクリプト群を、そのまま定期点検ツールとして利用する。"
        "「期待値(コード上の宣言)と実環境の状態を突合する」という設計のため、目視確認より確実に整合性を検知できる。"
    ))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["スクリプト", "点検内容", "推奨頻度"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    checks = [
        ("backend/verify_schema.py", "テーブル・列の定義が期待通りか", "月次、またはスキーマ変更時"),
        ("backend/verify_security_roles.py", "セキュリティロールの権限マトリクスが期待通りか", "月次、またはロール変更時"),
        ("map/verify_pcf_control.py", "地図コントロールがDataverseに正しく登録されているか", "地図コントロールのデプロイ時"),
        ("map/diagnose_map_binding.py", "地図コントロールがビュー/フォームに正しくバインドされているか", "地図表示に異常を感じた時"),
        ("map/check_env_var_value.py", "Azure Mapsサブスクリプションキーが有効か", "地図が表示されない時、キーローテーション時"),
        ("backend/check_audit_settings.py", "組織・テーブルの監査ログが有効なままか", "四半期ごと"),
        ("backend/run_load_test.py", "同時アクセス時の応答速度に劣化がないか", "データ量が大きく増えた時"),
    ]
    for row_data in checks:
        row = table.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = v

    add_h(doc, "3. 障害対応手順")
    add_p(doc, "障害検知から復旧までの標準フローを以下に示す。")
    steps = [
        "1. 検知: ユーザー報告、または定期点検スクリプトのエラーで異常を把握する",
        "2. 一次切り分け: 「デプロイ・実装が原因」か「Dataverse/Azure Maps側の障害」かを切り分ける。"
        "本プロジェクトの実例では、まずブラウザのコンソールログ(F12)を確認し、エラーメッセージの有無で判断した",
        "3. 詳細調査: 憶測で修正せず、関連する検証スクリプト(第2節)や直接のAPIクエリで実際の状態を確認する。"
        "例: 「地図が表示されない」場合、diagnose_map_binding.py → check_env_var_value.py → ブラウザコンソール、の順に切り分けた",
        "4. 復旧対応: 原因が特定できた修正を実施し、再度該当する検証スクリプトで解消を確認する",
        "5. 記録: 発生日時・症状・原因・対応内容を記録する(第4節のフォーマットを使用)",
    ]
    for s in steps:
        add_p(doc, s)

    add_h(doc, "4. インシデント記録フォーマット")
    add_p(doc, "セキュリティ関連を含むすべてのインシデントは、以下の項目で記録する。")
    add_bullets(doc, [
        "発生日時",
        "検知経路(ユーザー報告/定期点検/監視アラート)",
        "症状(何が起きたか)",
        "影響範囲(対象データ・対象ユーザー)",
        "原因(根本原因。憶測ではなく確認できた事実を記載)",
        "対応内容・復旧日時",
        "再発防止策",
    ])

    add_h(doc, "5. 再発防止プロセス")
    add_p(doc, (
        "本プロジェクトの開発過程で、実際に同種の不具合(PCFコントロールの「デプロイは成功するが動作しない」系の"
        "問題)を複数回経験した。1件目(地図タイル非表示)の原因調査で得た知見(ビルドツールのnode_modules誤"
        "トランスパイル)は、fix-pcf-webpack.pyとしてハーネス化し、postinstallフックで自動的に再発を防止する"
        "仕組みに変換した。このように、再発防止策は「文書に書くだけ」ではなく、可能な限り自動化・ハーネス化する"
        "方針を採る。"
    ))

    add_h(doc, "6. エスカレーション")
    add_p(doc, (
        "個人ポートフォリオのため、実際の組織的エスカレーション体制はない。実運用を想定する場合は、"
        "「一次対応者(運用担当)→システム管理者→委託業者」の3段階を想定し、各段階の対応可能時間・"
        "連絡手段を別途定める必要がある。"
    ))

    disclaimer = doc.add_paragraph(
        "本プロジェクトは東京都港区が公表した公募資料を参考にした個人の学習・ポートフォリオ目的の制作物であり、"
        "港区への正式な提案書・提出物ではありません。"
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.save(OUT_FILE)
    print(f"運用・障害対応手順書.docx を生成しました: {OUT_FILE}")


if __name__ == "__main__":
    main()
