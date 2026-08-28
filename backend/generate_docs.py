"""
generate_docs.py
------------------
(仮称)樹木管理システム ポートフォリオ用 ドキュメント生成ハーネス

以下を「単一の情報源」として、README.md と 一枚物資料(one_pager.docx) を自動生成する。
    - schema_manifest.json        … データモデル(テーブル・列)の実態
    - create_security_roles.py    … セキュリティロールの権限マトリクス設計
    - project_meta.json           … 人が書く説明文(概要・ハイライト等)

テーブルやロールを追加したら、このスクリプトを再実行するだけで
ドキュメントが最新の状態に更新される。

事前準備:
    pip install python-docx

実行方法:
    py generate_docs.py
"""

import json
import os
import sys

from dataverse_common import load_manifest
from create_security_roles import ENTITIES, ROLES, MATRIX, DEPTH_MAP, CRUD_TO_TYPE

HERE = os.path.dirname(os.path.abspath(__file__))
META_FILE = os.path.join(HERE, "project_meta.json")
README_FILE = os.path.join(HERE, "README.md")
ONEPAGER_FILE = os.path.join(HERE, "one_pager.docx")

DEPTH_LABEL_JP = {"Global": "組織全体", "Local": "担当エリア(BU)", "Basic": "自分のみ"}


def load_meta():
    with open(META_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def compute_stats(manifest):
    table_count = len(manifest)
    column_count = sum(len(info["columns"]) for info in manifest.values())
    role_count = len(ROLES) - 1  # システム管理者(組み込み)を除く
    perm_count = 0
    for entity_suffix, role_map in MATRIX.items():
        for role_name, (crud, scope) in role_map.items():
            if role_name == "システム管理者" or not crud or scope == "なし":
                continue
            types_needed = len(crud)
            if "C" in crud or "U" in crud:
                types_needed += 2  # Append, AppendTo
            perm_count += types_needed
    return {
        "table_count": table_count,
        "column_count": column_count,
        "role_count": role_count,
        "perm_count": perm_count,
    }


def build_mermaid_er(manifest, relationships):
    lines = ["```mermaid", "erDiagram"]
    for rel in relationships:
        frm = rel["from"].upper()
        to = rel["to"].upper()
        lines.append(f'    {frm} ||--o{{ {to} : "{rel["label"]}"')
    lines.append("```")
    legend = ["", "凡例(テーブル論理名 → 表示名):"]
    for suffix, info in manifest.items():
        legend.append(f"- `{suffix.upper()}` = {info['display']}")
    return "\n".join(lines + legend)


def build_table_section(manifest):
    parts = []
    for suffix, info in manifest.items():
        parts.append(f"### {info['display']} (`{suffix}`)\n")
        parts.append(f"列数: {len(info['columns'])}件\n")
        parts.append("| # | 列名 |")
        parts.append("|---|---|")
        for i, col in enumerate(info["columns"], start=1):
            parts.append(f"| {i} | {col['display']} |")
        parts.append("")
    return "\n".join(parts)


def build_progress_section(meta):
    prog = meta.get("requirement_progress")
    if not prog:
        return ""
    f = prog["functional"]
    nf = prog["non_functional"]
    lines = [
        f"({prog['as_of']}時点、コードベースとの突き合わせによる評価)",
        "",
        "### 機能要件(35項目)",
        f"- 実装済み(標準機能で充足含む): {f['done']}件",
        f"- 一部実装/仕様と差異あり: {f['partial']}件",
        f"- 未着手: {f['not_started']}件",
        f"- 要確認(コードからは判断不可): {f['unconfirmed']}件",
        "",
        "### 非機能要件(44項目)",
        f"- 実装済み(標準機能で充足含む): {nf['done']}件",
        f"- 未着手(主に文書系成果物): {nf['not_started']}件",
        f"- 要確認(設定状況等、コードからは判断不可): {nf['unconfirmed']}件",
        "",
        "### 主な未実装・要確認事項",
    ]
    lines += [f"- {g}" for g in prog["notable_gaps"]]
    lines.append("")
    lines.append(f"項目別の判定根拠: {prog['detail_source']}")
    return "\n".join(lines)


def build_role_matrix_section():
    header = "| テーブル \\ ロール | " + " | ".join(r for r in ROLES if r != "システム管理者") + " |"
    sep = "|---" * (len(ROLES)) + "|"
    lines = [header, sep]
    for entity_suffix, role_map in MATRIX.items():
        row = [entity_suffix]
        for role in ROLES:
            if role == "システム管理者":
                continue
            crud, scope = role_map.get(role, ("", "なし"))
            if not crud or scope == "なし":
                row.append("―")
            else:
                depth_jp = DEPTH_LABEL_JP.get(DEPTH_MAP[scope], scope)
                row.append(f"{crud}({depth_jp})")
        lines.append("| " + " | ".join(row) + " |")
    note = (
        "\n凡例: C=作成 R=参照 U=更新 D=削除。"
        "「システム管理者」はDataverse組み込みロールで全テーブルにフルアクセスのため表から省略。"
    )
    return "\n".join(lines) + note


def generate_readme(meta, manifest, stats):
    er_diagram = build_mermaid_er(manifest, meta["relationships"])
    table_section = build_table_section(manifest)
    role_section = build_role_matrix_section()
    progress_section = build_progress_section(meta)

    content = f"""# {meta['title']}

> {meta['tagline']}

## 概要

{meta['purpose']}

## 主な実績(自動集計)

- テーブル数: {stats['table_count']}
- 列数(合計): {stats['column_count']}
- セキュリティロール数: {stats['role_count']}(＋組み込みのシステム管理者)
- 付与した権限数(合計): {stats['perm_count']}

## 技術スタック

{chr(10).join('- ' + t for t in meta['tech_stack'])}

## データモデル(ER図)

{er_diagram}

## テーブル定義

{table_section}

## セキュリティロール権限マトリクス

{role_section}

## ハーネス設計

### 構築ハーネス
{meta['harness_design']['build']}

### 検証ハーネス
{meta['harness_design']['verify']}

### ドキュメント生成ハーネス
{meta['harness_design']['docs']}

## 実装のハイライト

{chr(10).join(f"### {h['title']}{chr(10)}{h['body']}{chr(10)}" for h in meta['highlights'])}

## 仕様書要件に対する実装進捗

{progress_section}

## セットアップ手順

{chr(10).join(f"{i+1}. {s}" for i, s in enumerate(meta['setup_steps']))}

---

{meta['disclaimer']}

*このREADMEは generate_docs.py によって自動生成されています。手動で編集した内容は次回実行時に上書きされます。*
"""
    with open(README_FILE, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"README.md を生成しました: {README_FILE}")


def generate_onepager(meta, stats):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx がインストールされていません。'pip install python-docx' を実行してください。")
        return

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = doc.add_heading(meta["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tagline = doc.add_paragraph(meta["tagline"])
    tagline.runs[0].italic = True
    tagline.runs[0].font.size = Pt(11)

    doc.add_paragraph(meta["purpose"])

    doc.add_heading("主な実績", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    labels = ["テーブル数", "列数", "セキュリティロール数", "付与権限数"]
    values = [stats["table_count"], stats["column_count"], stats["role_count"], stats["perm_count"]]
    for i, (label, value) in enumerate(zip(labels, values)):
        hdr[i].text = f"{label}\n{value}"

    doc.add_heading("技術スタック", level=2)
    for t in meta["tech_stack"]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("実装のハイライト", level=2)
    for h in meta["highlights"]:
        p = doc.add_paragraph()
        run = p.add_run(h["title"])
        run.bold = True
        doc.add_paragraph(h["body"])

    disclaimer = doc.add_paragraph(meta["disclaimer"])
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.save(ONEPAGER_FILE)
    print(f"one_pager.docx を生成しました: {ONEPAGER_FILE}")


def main():
    manifest = load_manifest()
    if not manifest:
        print("schema_manifest.json が見つからないか空です。先にテーブル作成スクリプトを実行してください。")
        sys.exit(1)

    meta = load_meta()
    stats = compute_stats(manifest)

    generate_readme(meta, manifest, stats)
    generate_onepager(meta, stats)

    print("\n完了しました。README.md と one_pager.docx を確認してください。")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # noqa: BLE001
        print(f"\nエラーが発生しました: {exc}", file=sys.stderr)
        sys.exit(1)
