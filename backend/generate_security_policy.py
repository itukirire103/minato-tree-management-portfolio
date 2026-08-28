# -*- coding: utf-8 -*-
"""
generate_security_policy.py
------------------------------
(仮称)樹木管理システム ポートフォリオ用
セキュリティ設計方針書(.docx)を生成する。

非機能要件#24(総務省ガイドライン等を参考にした設計方針の文書化)に対応。
港区の内部規定は入手できないため、総務省「地方公共団体における情報セキュリティ
ポリシーに関するガイドライン」の一般的な構成(組織的・人的・物理的・技術的対策)
を参考に、本システムで実際に実施した対策を整理している。

実行方法:
    py generate_security_policy.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
import os

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_FILE = os.path.join(HERE, "セキュリティ設計方針書.docx")


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("セキュリティ設計方針書", level=0)
    sub = doc.add_paragraph("(仮称)樹木管理システム 構築プロジェクト")
    sub.runs[0].italic = True
    doc.add_paragraph("作成日: 2026-08-28")

    doc.add_heading("1. 目的", level=1)
    doc.add_paragraph(
        "本システムが取り扱う樹木情報・業務データを、不正アクセス・改ざん・漏えいから保護するための"
        "セキュリティ対策方針を定める。"
    )

    doc.add_heading("2. 参照ガイドライン", level=1)
    doc.add_paragraph(
        "東京都港区の内部セキュリティ規定は公開情報として入手できないため、総務省が公表している"
        "「地方公共団体における情報セキュリティポリシーに関するガイドライン」の一般的な構成"
        "(組織的対策・人的対策・物理的対策・技術的対策の4分類)を参考に、本システムで実施した"
        "対策を整理する。"
    )

    doc.add_heading("3. 組織的対策", level=1)
    doc.add_paragraph("システム管理者ロールを設け、権限の付与・変更を一元管理する体制とした。", style="List Bullet")
    doc.add_paragraph(
        "セキュリティロールは業務上必要な最小限の権限(最小権限の原則)に基づき、6区分"
        "(各所管理者/区一般職員/街路樹管理委託事業者/協定管理者/その他閲覧専用/システム管理者)"
        "に分けて設計した。", style="List Bullet")

    doc.add_heading("4. 人的対策", level=1)
    doc.add_paragraph("利用者アカウントは組織アカウントに限定し、個人所有のアカウントでのアクセスを許可しない。", style="List Bullet")
    doc.add_paragraph("委託業者・協定管理者は担当エリア外のデータにアクセスできないよう権限を制限した。", style="List Bullet")

    doc.add_heading("5. 物理的対策", level=1)
    doc.add_paragraph(
        "Power Platform環境の作成時にリージョン(データセンター所在地)として「日本」を選択し、"
        "国内データセンターでのデータ保管を実現している。物理的なデータセンターの管理体制"
        "(入退室管理・機器管理等)はMicrosoft側の責任範囲であり、ISO27001等の第三者認証を"
        "根拠として妥当性を確認している。", style="List Bullet")

    doc.add_heading("6. 技術的対策(実装状況)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["対策項目", "内容", "実装状況"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    items = [
        ("認証", "多要素認証(MFA)を必須化", "実装済み(セキュリティの既定値群で有効化を確認)"),
        ("認可", "セキュリティロールによるアクセス制御(6区分、テーブル・フィールド単位)", "実装済み(verify_security_roles.pyで自動検証)"),
        ("監査", "組織全体・テーブル単位の変更履歴の記録", "実装済み(組織全体はWeb API、テーブル単位は管理画面から有効化)"),
        ("データ保護", "国内データセンターでの保管、通信の暗号化(HTTPS)", "実装済み(Power Platform標準機能)"),
        ("論理削除", "削除データを履歴として保持するソフトデリート", "Dataverse標準機能として利用可能。本プロジェクトでの動作確認は未実施"),
        ("パッチ管理", "セキュリティパッチの適用", "SaaSのためMicrosoft側が自動実施"),
        ("バックアップ", "定期的なデータバックアップと復元手順の確立", "手動バックアップの作成・復元コマンドの動作を確認済み(2026-08-28)。自動日次バックアップは本番相当環境が必要"),
        ("負荷対策", "同時アクセス時の応答性能の確認", "疑似10並列アクセスで実測済み(平均応答624ms)"),
    ]
    for row_data in items:
        row = table.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = v

    doc.add_heading("7. 運用ルール", level=1)
    doc.add_paragraph("アカウントの棚卸し(異動・退職者の権限削除)を定期的に実施する。", style="List Bullet")
    doc.add_paragraph("インシデント発生時は「運用・障害対応手順書」の記録フォーマットに従い記録・共有する。", style="List Bullet")
    doc.add_paragraph("定期点検スクリプト(verify_schema.py等)を用いて、意図しない設定変更がないか確認する。", style="List Bullet")

    doc.add_heading("8. 今後の見直し", level=1)
    doc.add_paragraph(
        "本方針は個人ポートフォリオとしての設計方針であり、実際の自治体調達においては、当該自治体の"
        "情報セキュリティポリシーに照らした個別の見直し・承認プロセスが必要になる。"
    )

    disclaimer = doc.add_paragraph(
        "本プロジェクトは東京都港区が公表した公募資料を参考にした個人の学習・ポートフォリオ目的の制作物であり、"
        "港区への正式な提案書・提出物ではありません。"
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.save(OUT_FILE)
    print(f"セキュリティ設計方針書.docx を生成しました: {OUT_FILE}")


if __name__ == "__main__":
    main()
