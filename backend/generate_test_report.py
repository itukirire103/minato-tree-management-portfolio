# -*- coding: utf-8 -*-
"""
generate_test_report.py
--------------------------
(仮称)樹木管理システム ポートフォリオ用
テスト計画書・実施結果報告書(.docx)を生成する。

非機能要件#41(テスト計画書・テストケース)・#43(実施計画書・結果報告書)に対応。
2026-08-28の開発・デバッグセッションで実際に発見・修正した不具合と、
実際に取得した負荷テスト結果を、そのままテストケースの実施結果として記録している
(架空のテストではなく、実際の開発過程の記録)。

実行方法:
    pip install python-docx
    py generate_test_report.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_FILE = os.path.join(HERE, "テスト計画書_実施結果報告書.docx")

TEST_CASES = [
    # (ID, 分類, テスト項目, 内容, 期待結果, 実施結果, 備考)
    ("A1", "バックエンド基盤", "テーブル・列のスキーマ検証",
     "schema_manifest.jsonの自己申告内容とDataverse実環境の状態を突合する",
     "全テーブル・全列が一致する", "合格",
     "verify_schema.pyによる自動検証"),
    ("A2", "バックエンド基盤", "セキュリティロール権限検証",
     "create_security_roles.pyの権限マトリクスとDataverse実環境の権限設定を突合する",
     "6ロール・95権限が期待通りに設定されている", "合格",
     "verify_security_roles.pyによる自動検証"),
    ("A3", "バックエンド基盤", "業務ロジック(健全度自動更新)",
     "樹木診断結果を登録すると、樹木マスタの健全度が自動更新される",
     "健全度が診断結果と一致して更新される", "合格",
     "Power Automateクラウドフロー。初回実装時に画面反映されない事象があり、原因(APIで作成した列がフォームに未反映)を特定して解消"),
    ("A4", "バックエンド基盤", "業務ロジック(植替えステータス自動更新)",
     "植替え履歴を登録すると、旧樹木のステータスが「植替え済」に自動更新される",
     "旧樹木のステータスが自動更新される", "合格",
     "A3の知見を活かし、事前準備を整えた上で1回のテストで合格"),
    ("B1", "地図機能", "地図タイル表示",
     "「樹木マップ」ビューを開いた際、Azure Mapsの地図タイルが表示される",
     "地図タイルが表示される", "修正後に合格",
     "初回はfeature-usage宣言不備、既定ビュー未選択、pcf-scriptsによるnode_modules誤トランスパイルの3段階の要因で失敗。原因を1つずつ切り分けて解消"),
    ("B2", "地図機能", "樹木ピンの色分け表示",
     "健全度(A/B1/B2/C)に応じてピンの色が変わる",
     "健全度に応じた色でピンが表示される", "合格", ""),
    ("B3", "地図機能", "全件表示の確認",
     "投入した約300件の樹木データが、初回表示時点で全件地図に表示される",
     "全件が表示される", "修正後に合格",
     "初回はpaging.setPageSize(500)が初回ロードに間に合わず、一部しか表示されない不具合があった。loadNextPage()で全件読み込む実装に変更して解消"),
    ("B4", "地図機能", "ポイントクリックでの詳細画面遷移",
     "地図上の樹木ピンをクリックすると、該当レコードの詳細フォームが開く",
     "詳細フォームが開く", "合格", "機能要件#28"),
    ("B5", "地図機能", "ドラッグによる位置修正",
     "更新権限を持つユーザーが樹木ピンをドラッグすると、緯度経度が更新される",
     "ドラッグ先の座標でDataverseのレコードが更新される", "修正後に合格",
     "機能要件#30/#31。初回はドラッグ操作が地図全体のパン操作と競合し、ポイントではなく地図全体が動く不具合があった"),
    ("B6", "地図機能", "空き地クリックでの新規登録",
     "作成権限を持つユーザーが地図上の空き地をクリックすると、その位置を初期値にした新規登録が行える",
     "新規レコードが正しいステータス(現存)・緯度経度で作成される", "修正後に合格",
     "機能要件#29。openForm()のparameters引数は選択肢型フィールドの既定値を反映しないことが判明し、webAPI.createRecord()で確実に作成する方式に変更"),
    ("C1", "非機能要件", "同時アクセス負荷テスト",
     "10並列相当の同時アクセスをDataverse Web APIに対して実施する",
     "エラーなく応答が返る", "合格",
     "run_load_test.py実施。100リクエスト全て成功、平均応答624ms、95パーセンタイル728ms(2026-08-28実測)"),
    ("C2", "非機能要件", "検索速度",
     "約300件規模のデータに対する一覧取得の応答速度を計測する",
     "体感できる遅延がない", "合格",
     "C1と同時に計測。約300件規模では良好。5,200件相当の規模は未検証"),
    ("C3", "非機能要件", "多要素認証(MFA)",
     "Microsoft Entra IDでMFAが有効になっているか確認する",
     "MFAが有効", "合格",
     "「セキュリティの既定値群」が既定で有効になっていることを確認"),
    ("C4", "非機能要件", "監査ログ",
     "組織全体・樹木マスタテーブルの監査ログが有効になっているか確認する",
     "両方とも有効", "合格",
     "組織全体はWeb APIで、テーブル単位はPower Apps管理画面から有効化し、APIで確認"),
    ("C5", "非機能要件", "バックアップ作成・復元",
     "手動バックアップを作成し、復元操作を実施する",
     "バックアップが作成され、復元が完了する", "部分合格",
     "バックアップ作成・存在確認・復元コマンドの実行までは確認。作成直後のバックアップは復元可能になるまで時間を要する仕様のため、復元完了までは未実施"),
    ("C6", "非機能要件", "ポイントインタイム復元",
     "特定時点への復元機能が利用可能か確認する",
     "-", "対象外(利用不可)",
     "環境種別がDeveloperであるため、本番相当環境限定のこの機能は利用不可であることを確認(pac admin list)"),
    ("C7", "非機能要件", "作業前後写真の添付",
     "作業履歴レコードのタイムラインに画像ファイルを添付できるか確認する",
     "画像ファイルを添付・プレビューできる", "合格",
     "Dataverse標準のノート機能(HasNotes)で実現。専用列の追加は不要と判断"),
]

SUMMARY_COUNTS = {
    "合格": sum(1 for c in TEST_CASES if c[5] == "合格"),
    "修正後に合格": sum(1 for c in TEST_CASES if c[5] == "修正後に合格"),
    "部分合格": sum(1 for c in TEST_CASES if c[5] == "部分合格"),
    "対象外(利用不可)": sum(1 for c in TEST_CASES if c[5] == "対象外(利用不可)"),
}


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("テスト計画書・実施結果報告書", level=0)
    sub = doc.add_paragraph("(仮称)樹木管理システム 構築プロジェクト")
    sub.runs[0].italic = True

    doc.add_paragraph("実施日: 2026-08-28")
    doc.add_paragraph("対象システム: (仮称)樹木管理システム(Power Platform/Dataverse + Azure Maps PCF)")

    doc.add_heading("1. テスト方針", level=1)
    doc.add_paragraph(
        "本プロジェクトは個人ポートフォリオのため、実利用者の代わりに開発者自身がテスト実施者を兼ねる。"
        "テストは大きく3種類に分類する:"
    )
    doc.add_paragraph("バックエンド基盤(Dataverseのテーブル・セキュリティロール・業務ロジック)の検証", style="List Bullet")
    doc.add_paragraph("地図機能(Azure Maps PCFコントロール)の機能検証", style="List Bullet")
    doc.add_paragraph("非機能要件(性能・セキュリティ・可用性)の検証", style="List Bullet")
    doc.add_paragraph(
        "重要な方針として、不具合が発生した際は憶測で修正せず、Dataverse Web APIへの直接クエリ・"
        "ブラウザコンソールログの精読・実データの直接確認など、根拠に基づいて原因を切り分けてから対処する。"
        "以下のテストケースの「備考」欄には、実際の開発過程で発生した不具合とその原因・対処を記録している。"
    )

    doc.add_heading("2. テストケース一覧・実施結果", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "分類", "テスト項目", "内容", "期待結果", "結果", "備考"]
    widths = [0.6, 1.6, 2.2, 3.5, 2.8, 1.5, 5.0]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for case in TEST_CASES:
        row = table.add_row().cells
        for i, value in enumerate(case):
            set_cell_text(row[i], value)

    doc.add_heading("3. 結果サマリー", level=1)
    for status, count in SUMMARY_COUNTS.items():
        doc.add_paragraph(f"{status}: {count}件", style="List Bullet")
    doc.add_paragraph(f"総テストケース数: {len(TEST_CASES)}件")

    doc.add_heading("4. 総括", level=1)
    doc.add_paragraph(
        "地図機能(B1, B3, B5, B6)で複数の不具合を発見・修正した。いずれも「デプロイ・実装は完了しているが"
        "実際には動作しない」という種類の不具合であり、コンソールログの読解とDataverse/Azure Maps各APIへの"
        "直接的な問い合わせを組み合わせることで、憶測に頼らず原因を特定できた。"
    )
    doc.add_paragraph(
        "非機能要件のうち、実運用に必要な設定作業(MFA・監査ログ)は既に有効、または今回のセッションで"
        "有効化・確認した。バックアップの完全な復元完了とポイントインタイム復元は、Developer環境の制約により"
        "本番相当環境でなければ検証できないことを確認した。これは個人ポートフォリオとしての環境制約であり、"
        "実際の自治体調達においては本番相当のSandbox/Production環境で追加検証が必要になる点として明記する。"
    )

    disclaimer = doc.add_paragraph(
        "本プロジェクトは東京都港区が公表した公募資料を参考にした個人の学習・ポートフォリオ目的の制作物であり、"
        "港区への正式な提案書・提出物ではありません。"
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.save(OUT_FILE)
    print(f"テスト計画書・実施結果報告書.docx を生成しました: {OUT_FILE}")


if __name__ == "__main__":
    main()
